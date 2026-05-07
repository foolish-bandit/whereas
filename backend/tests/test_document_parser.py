"""Tests for the document parser.

Fixtures are generated in-test (reportlab for PDFs, python-docx for DOCXs)
so the contents of the synthetic NDA used in every test are visible right
here in code, and no binary files end up in the repo.
"""
from __future__ import annotations

import hashlib
import io
import sys

import pytest
from docx import Document

reportlab = pytest.importorskip("reportlab")
from reportlab.lib.pagesizes import letter  # noqa: E402
from reportlab.pdfgen import canvas  # noqa: E402

from app.services.document_parser import (  # noqa: E402
    DocumentParseError,
    DocumentParseTimeoutError,
    DocumentTooLargeError,
    ParsedDocument,
    UnsupportedDocumentTypeError,
    _run_in_subprocess,
    parse_document,
)

# A synthetic, public-domain NDA used as the source content for every fixture.
# These lines are deliberately not real client data.
_NDA_LINES: tuple[str, ...] = (
    "NON-DISCLOSURE AGREEMENT",
    "This Agreement is entered into between Acme, Inc. (Disclosing Party) "
    "and Beta LLC (Receiving Party).",
    "Effective Date: January 1, 2025.",
    "Term: Two (2) years from the Effective Date.",
    "Governing Law: This Agreement shall be governed by the laws of the "
    "State of Delaware.",
    "Confidential Information means any non-public information disclosed "
    "by the Disclosing Party.",
)


def _build_pdf(lines_per_page: list[list[str]]) -> bytes:
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    for page_lines in lines_per_page:
        y = 720
        for line in page_lines:
            c.drawString(72, y, line)
            y -= 24
        c.showPage()
    c.save()
    return buf.getvalue()


@pytest.fixture
def synthetic_pdf_bytes() -> bytes:
    """Two-page synthetic NDA PDF generated with reportlab."""
    return _build_pdf([list(_NDA_LINES[:3]), list(_NDA_LINES[3:])])


@pytest.fixture
def synthetic_docx_bytes() -> bytes:
    """Synthetic NDA DOCX generated with python-docx."""
    doc = Document()
    doc.add_heading(_NDA_LINES[0], level=1)
    for line in _NDA_LINES[1:]:
        doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _is_offline_model_error(err: BaseException) -> bool:
    """Detect a Docling failure caused by missing HuggingFace model downloads.

    Self-hosted deployments are expected to either have internet access on
    first run or pre-bake the layout/OCR models into the image. CI sandboxes
    that block huggingface.co will trip this branch and we skip the test
    rather than report a false failure.
    """
    msg = str(err).lower()
    needles = (
        "huggingface",
        "localentrynotfound",
        "failed to download",
        "host not in allowlist",
        "modelscope",
    )
    return any(n in msg for n in needles)


def _parse_or_skip(
    file_bytes: bytes, filename: str, **kwargs: object
) -> ParsedDocument:
    try:
        return parse_document(file_bytes, filename, **kwargs)  # type: ignore[arg-type]
    except DocumentParseTimeoutError as e:
        pytest.skip(
            "Docling/RapidOCR model startup exceeded timeout in this environment: "
            f"{e}"
        )
    except DocumentParseError as e:
        if _is_offline_model_error(e):
            pytest.skip(
                f"Docling model download unavailable in this environment: {e}"
            )
        raise


class TestParseOrSkip:
    def test_skips_docling_startup_timeout(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        def fake_parse_document(*_args: object, **_kwargs: object) -> ParsedDocument:
            raise DocumentParseTimeoutError("Parsing exceeded 180s budget.")

        monkeypatch.setattr(
            sys.modules[__name__], "parse_document", fake_parse_document
        )

        with pytest.raises(
            pytest.skip.Exception,
            match="Docling/RapidOCR model startup exceeded timeout",
        ):
            _parse_or_skip(b"%PDF", "nda.pdf", timeout_seconds=180)

    def test_skips_known_offline_model_download_error(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        def fake_parse_document(*_args: object, **_kwargs: object) -> ParsedDocument:
            raise DocumentParseError("HuggingFace model download failed")

        monkeypatch.setattr(
            sys.modules[__name__], "parse_document", fake_parse_document
        )

        with pytest.raises(
            pytest.skip.Exception,
            match="Docling model download unavailable",
        ):
            _parse_or_skip(b"%PDF", "nda.pdf", timeout_seconds=180)

    def test_reraises_ordinary_parse_error(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        def fake_parse_document(*_args: object, **_kwargs: object) -> ParsedDocument:
            raise DocumentParseError("ordinary parser regression")

        monkeypatch.setattr(
            sys.modules[__name__], "parse_document", fake_parse_document
        )

        with pytest.raises(DocumentParseError, match="ordinary parser regression"):
            _parse_or_skip(b"%PDF", "nda.pdf", timeout_seconds=180)


def _assert_offset_invariants(doc: ParsedDocument) -> None:
    """The load-bearing invariant the extraction layer relies on."""
    for page in doc.pages:
        assert doc.full_text[page.char_start : page.char_end] == page.text
        for block in page.blocks:
            assert (
                doc.full_text[block.char_start : block.char_end] == block.text
            )
            assert block.page_number == page.page_number


class TestPdfParsing:
    def test_parses_pdf_and_offsets_are_consistent(
        self, synthetic_pdf_bytes: bytes
    ) -> None:
        result = _parse_or_skip(
            synthetic_pdf_bytes, "nda.pdf", timeout_seconds=180
        )
        _assert_offset_invariants(result)
        assert result.page_count >= 1
        assert result.content_hash == hashlib.sha256(
            synthetic_pdf_bytes
        ).hexdigest()
        # Some recognisable content from the synthetic NDA should appear.
        assert "Effective Date" in result.full_text
        assert "Delaware" in result.full_text


class TestDocxParsing:
    def test_parses_docx_and_offsets_are_consistent(
        self, synthetic_docx_bytes: bytes
    ) -> None:
        result = _parse_or_skip(
            synthetic_docx_bytes, "nda.docx", timeout_seconds=180
        )
        _assert_offset_invariants(result)
        # DOCX has no native page concept; we synthesize a single page.
        assert result.page_count == 1
        assert len(result.pages) == 1
        assert "NON-DISCLOSURE AGREEMENT" in result.full_text
        assert "Delaware" in result.full_text
        # Bounding boxes are not surfaced by Docling for DOCX.
        for block in result.pages[0].blocks:
            assert block.bbox is None


class TestUnsupportedExtension:
    @pytest.mark.parametrize(
        "filename",
        [
            "spec.txt",
            "scan.png",
            "legacy.doc",
            "macro.docm",
            "noext",
            "archive.tar.gz",
        ],
    )
    def test_rejects_unsupported_extension(self, filename: str) -> None:
        with pytest.raises(UnsupportedDocumentTypeError) as exc_info:
            parse_document(b"irrelevant bytes", filename)
        msg = str(exc_info.value).lower()
        # Error text should tell the user what we accept.
        assert ".pdf" in msg
        assert ".docx" in msg

    def test_legacy_doc_message_suggests_resave(self) -> None:
        with pytest.raises(UnsupportedDocumentTypeError) as exc_info:
            parse_document(b"\xd0\xcf\x11\xe0", "legacy.doc")
        assert "docx" in str(exc_info.value).lower()


class TestHashDeterminism:
    def test_same_bytes_same_hash(self, synthetic_docx_bytes: bytes) -> None:
        a = _parse_or_skip(
            synthetic_docx_bytes, "first.docx", timeout_seconds=180
        )
        b = _parse_or_skip(
            synthetic_docx_bytes, "second.docx", timeout_seconds=180
        )
        assert a.content_hash == b.content_hash
        assert (
            a.content_hash
            == hashlib.sha256(synthetic_docx_bytes).hexdigest()
        )

    def test_one_bit_flip_produces_different_hash(
        self, synthetic_docx_bytes: bytes
    ) -> None:
        flipped = bytearray(synthetic_docx_bytes)
        flipped[-1] ^= 0x01
        a = _parse_or_skip(
            synthetic_docx_bytes, "original.docx", timeout_seconds=180
        )
        # The mutated bytes likely break the docx zip; we only assert the
        # *hash* differs, which is the point of the determinism test.
        assert a.content_hash != hashlib.sha256(bytes(flipped)).hexdigest()


def _slow_worker(conn, _file_bytes, _ext):  # pragma: no cover - runs in worker
    """Test-only worker that sleeps long enough to trip the timeout.

    Must be a top-level function so it's picklable for spawn.
    """
    import time

    time.sleep(30)
    conn.send(("ok", None))


class TestTimeout:
    def test_timeout_raises_and_kills_worker(self) -> None:
        with pytest.raises(DocumentParseTimeoutError):
            _run_in_subprocess(
                b"\x00",
                ".pdf",
                timeout_seconds=1,
                worker=_slow_worker,
            )


class TestMaxPagesGuard:
    def test_pdf_with_too_many_pages_rejected_pre_docling(self) -> None:
        # 5-page PDF, max_pages=2 → reject in the cheap pre-pass before
        # Docling is ever invoked.
        pdf_bytes = _build_pdf([[f"Page {i + 1}"] for i in range(5)])
        with pytest.raises(DocumentTooLargeError) as exc_info:
            parse_document(pdf_bytes, "huge.pdf", max_pages=2)
        msg = str(exc_info.value)
        assert "5" in msg and "2" in msg

    def test_pdf_within_limit_proceeds(
        self, synthetic_pdf_bytes: bytes
    ) -> None:
        # Two-page PDF with max_pages=10 should not trip the guard. The
        # parse may still skip on missing models; that's OK.
        result = _parse_or_skip(
            synthetic_pdf_bytes, "nda.pdf", timeout_seconds=180, max_pages=10
        )
        assert result.page_count >= 1
