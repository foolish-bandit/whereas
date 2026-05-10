"""API tests for the request -> contract conversion endpoint.

The conversion route is the seam between PR #47's intake/work-queue
foundation (``ContractRequest`` + ``InboxItem``) and PR #42's template
generation service. The point of this test module is to pin down both
sides of that seam:

* state transitions on the request and inbox sides;
* clean reuse of the generation service (variable validation propagates,
  no parallel taxonomy gets invented);
* failure semantics — the request must not flip to ``completed`` if
  generation blows up.
"""
from __future__ import annotations

import io
import secrets
import subprocess
import uuid
from collections.abc import AsyncIterator, Iterator
from dataclasses import dataclass
from typing import Any

import httpx
import pytest
from docx import Document  # type: ignore[import-not-found]
from sqlalchemy import event, select
from sqlalchemy.ext.asyncio import (
    AsyncEngine,
    AsyncSession,
    async_sessionmaker,
    create_async_engine,
)

try:
    from testcontainers.postgres import PostgresContainer
except ImportError:  # pragma: no cover
    PostgresContainer = None  # type: ignore[assignment,misc]

from app.api import agreement_templates as agreement_templates_api  # noqa: E402
from app.api import requests as requests_api  # noqa: E402
from app.core.database import Base, get_db  # noqa: E402
from app.main import app  # noqa: E402
from app.models import (  # noqa: E402
    AgreementTemplate,
    AgreementTemplateArtifact,
    AgreementTemplateMarkdownSnapshot,
    AgreementTemplateVariable,
    ApprovalPolicy,
    ApprovalStep,
    ApprovalWorkflowRun,
    ApprovalWorkflowTemplate,
    ApprovalWorkflowTemplateStep,
    Contract,
    ContractArtifact,
    ContractMarkdownSnapshot,
    ContractRequest,
    InboxItem,
    Organization,
    User,
)
from app.security.audit_log import AuditEvent  # noqa: E402
from app.security.encryption import create_org_master_key  # noqa: E402
from app.services.document_markdown import MarkdownConversionResult  # noqa: E402
from app.services.storage import StoredDocument  # noqa: E402

_PG_IMAGE = "pgvector/pgvector:pg16"
_INSTANCE_KEY = secrets.token_bytes(32)
_DOCX_MIME = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


def _docker_available() -> bool:
    if PostgresContainer is None:
        return False
    try:
        result = subprocess.run(
            ["docker", "info"], capture_output=True, timeout=5, check=False
        )
    except (FileNotFoundError, subprocess.TimeoutExpired):
        return False
    return result.returncode == 0


def _container_async_url(container: Any) -> str:
    sync_url = container.get_connection_url()
    if sync_url.startswith("postgresql+psycopg2://"):
        return sync_url.replace(
            "postgresql+psycopg2://", "postgresql+asyncpg://", 1
        )
    if sync_url.startswith("postgresql://"):
        return sync_url.replace("postgresql://", "postgresql+asyncpg://", 1)
    return sync_url


@pytest.fixture(scope="module")
def postgres_container() -> Iterator[Any | None]:
    if not _docker_available() or PostgresContainer is None:
        yield None
        return
    container = PostgresContainer(_PG_IMAGE)
    container.start()
    try:
        yield container
    finally:
        container.stop()


@pytest.fixture
async def engine(postgres_container: Any | None) -> AsyncIterator[AsyncEngine]:
    if postgres_container is None:
        engine = create_async_engine(
            "sqlite+aiosqlite:///:memory:", echo=False
        )
        # Approval-policy tables are required because PATCHing a
        # request fields like ``linked_template_id`` re-runs
        # ``apply_approval_policies_to_request`` (PR #53).
        tables = [
            Organization.__table__,
            User.__table__,
            AuditEvent.__table__,
            Contract.__table__,
            ContractArtifact.__table__,
            ContractMarkdownSnapshot.__table__,
            AgreementTemplate.__table__,
            AgreementTemplateArtifact.__table__,
            AgreementTemplateMarkdownSnapshot.__table__,
            AgreementTemplateVariable.__table__,
            ContractRequest.__table__,
            InboxItem.__table__,
            ApprovalWorkflowTemplate.__table__,
            ApprovalWorkflowTemplateStep.__table__,
            ApprovalWorkflowRun.__table__,
            ApprovalStep.__table__,
            ApprovalPolicy.__table__,
        ]
    else:
        engine = create_async_engine(
            _container_async_url(postgres_container), echo=False
        )
        tables = list(Base.metadata.sorted_tables)

    if engine.dialect.name == "sqlite":
        from pgvector.sqlalchemy import Vector
        from sqlalchemy.ext.compiler import compiles

        @compiles(Vector, "sqlite")
        def _compile_vector_for_sqlite(_t: Any, _c: Any, **_kw: Any) -> str:
            return "BLOB"

        @event.listens_for(engine.sync_engine, "connect")
        def _enable_sqlite_fk(dbapi_connection: Any, _r: Any) -> None:
            cursor = dbapi_connection.cursor()
            cursor.execute("PRAGMA foreign_keys=ON")
            cursor.close()

    async with engine.begin() as conn:
        if engine.dialect.name == "postgresql":
            await conn.exec_driver_sql("CREATE EXTENSION IF NOT EXISTS vector")
        await conn.run_sync(Base.metadata.drop_all, tables=tables)
        await conn.run_sync(Base.metadata.create_all, tables=tables)
    try:
        yield engine
    finally:
        async with engine.begin() as conn:
            await conn.run_sync(Base.metadata.drop_all, tables=tables)
        await engine.dispose()


@pytest.fixture
async def db_session(engine: AsyncEngine) -> AsyncIterator[AsyncSession]:
    maker = async_sessionmaker(engine, expire_on_commit=False, autoflush=False)
    async with maker() as session:
        yield session


@dataclass
class UserOrg:
    org: Organization
    user: User


def _wrapped_org_key(org_id: uuid.UUID) -> bytes:
    return create_org_master_key(
        organization_id=str(org_id),
        instance_key=_INSTANCE_KEY,
    ).to_bytes()


async def _create_user_org(
    session: AsyncSession,
    *,
    email: str | None = None,
) -> UserOrg:
    org = Organization(
        id=uuid.uuid4(),
        name=f"Org {uuid.uuid4()}",
        wrapped_master_key=_wrapped_org_key(uuid.uuid4()),
    )
    org.wrapped_master_key = _wrapped_org_key(org.id)
    user = User(
        id=uuid.uuid4(),
        organization_id=org.id,
        email=email or f"{uuid.uuid4()}@example.com",
        password_hash="hash",
        display_name="Test User",
        is_active=True,
    )
    session.add_all([org, user])
    await session.commit()
    return UserOrg(org=org, user=user)


def _headers(user: User | uuid.UUID) -> dict[str, str]:
    """Build the dev-user header.

    Accepting either a ``User`` row or a raw ``UUID`` keeps test code
    short and dodges a SQLAlchemy footgun: when an endpoint returns a
    4xx, FastAPI propagates the ``HTTPException`` through the
    dependency-cleanup path which calls ``rollback()`` on our test
    session, expiring every attached ``User``. Subsequent attribute
    access on ``user_org.user`` would then trigger a re-load from a
    closed greenlet context. Tests that issue a follow-up call after a
    rolled-back request use the captured UUID directly.
    """
    user_id = user.id if isinstance(user, User) else user
    return {"X-Whereas-Dev-User": str(user_id)}


def _make_nda_template_docx() -> bytes:
    doc = Document()
    doc.add_heading("Mutual NDA", level=1)
    doc.add_paragraph(
        "This Mutual Non-Disclosure Agreement is entered into by "
        "{{counterparty_name}} as of {{effective_date}}."
    )
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


_NDA_BYTES = _make_nda_template_docx()


class FakeStorage:
    """In-memory ``DocumentStorage`` stub. Mirrors
    test_template_generation_api.FakeStorage to keep both surfaces
    exercising the same fake.
    """

    _blobs: dict[str, bytes] = {}

    def __init__(self, _settings: Any) -> None:
        pass

    @classmethod
    def reset(cls) -> None:
        cls._blobs.clear()

    async def store_encrypted(
        self,
        *,
        plaintext_bytes: bytes,
        document_id: str,
        org_master_key: bytes,
    ) -> StoredDocument:
        s3_key = f"documents/{document_id}.enc"
        FakeStorage._blobs[s3_key] = plaintext_bytes
        return StoredDocument(
            s3_key=s3_key,
            wrapped_dek_bytes=b"wrapped-dek-" + document_id.encode()[:16],
            encrypted_blob_sha256="a" * 64,
            size_bytes=len(plaintext_bytes),
        )

    async def retrieve_decrypted(
        self,
        *,
        s3_key: str,
        document_id: str,
        wrapped_dek_bytes: bytes,
        org_master_key: bytes,
        expected_blob_sha256: str | None = None,
    ) -> bytes:
        return FakeStorage._blobs[s3_key]


@pytest.fixture
async def client(
    db_session: AsyncSession, monkeypatch: pytest.MonkeyPatch
) -> AsyncIterator[httpx.AsyncClient]:
    monkeypatch.setenv("WHEREAS_INSTANCE_KEY", _INSTANCE_KEY.hex())
    FakeStorage.reset()

    async def override_get_db() -> AsyncIterator[AsyncSession]:
        try:
            yield db_session
            await db_session.commit()
        except Exception:
            await db_session.rollback()
            raise

    app.dependency_overrides[get_db] = override_get_db

    # Both the templates surface and the requests surface instantiate
    # ``DocumentStorage``; both have to point at the same in-memory blob
    # store so the round-trip works.
    monkeypatch.setattr(agreement_templates_api, "DocumentStorage", FakeStorage)
    monkeypatch.setattr(requests_api, "DocumentStorage", FakeStorage)

    def _ok_convert(
        *,
        file_bytes: bytes,
        mime_type: str,
        filename: str | None,
        fallback_plain_text: str | None,
    ) -> MarkdownConversionResult:
        return MarkdownConversionResult(
            status="ready",
            markdown_text="# Generated\n\nbody",
            converter_name="fake",
            converter_version="0.0.1",
            warnings=[],
        )

    monkeypatch.setattr(
        agreement_templates_api, "convert_document_to_markdown", _ok_convert
    )
    from app.services import document_markdown

    monkeypatch.setattr(
        document_markdown, "convert_document_to_markdown", _ok_convert
    )

    class _StubParsed:
        full_text = "Template plain text fallback"

    monkeypatch.setattr(
        agreement_templates_api,
        "parse_document",
        lambda file_bytes, filename: _StubParsed(),
    )

    transport = httpx.ASGITransport(app=app)
    async with httpx.AsyncClient(transport=transport, base_url="http://test") as c:
        yield c
    app.dependency_overrides.clear()


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


async def _create_template_with_docx(
    client: httpx.AsyncClient,
    user: User,
    *,
    name: str = "Mutual NDA",
) -> str:
    created = await client.post(
        "/api/agreement-templates",
        headers=_headers(user),
        json={"name": name, "template_type": "NDA"},
    )
    assert created.status_code == 201, created.text
    template_id = created.json()["id"]
    upload = await client.post(
        f"/api/agreement-templates/{template_id}/upload",
        headers=_headers(user),
        files={"file": ("nda.docx", _NDA_BYTES, _DOCX_MIME)},
    )
    assert upload.status_code == 201, upload.text
    return template_id


async def _add_required_text_var(
    client: httpx.AsyncClient,
    user: User,
    template_id: str,
    *,
    key: str,
    variable_type: str = "text",
    sort_order: int = 0,
) -> None:
    response = await client.post(
        f"/api/agreement-templates/{template_id}/variables",
        headers=_headers(user),
        json={
            "key": key,
            "label": key.replace("_", " ").title(),
            "variable_type": variable_type,
            "required": True,
            "sort_order": sort_order,
        },
    )
    assert response.status_code == 201, response.text


async def _create_request_with_template(
    client: httpx.AsyncClient,
    user: User,
    *,
    template_id: str | None,
    title: str = "NDA with Acme",
) -> dict[str, Any]:
    payload: dict[str, Any] = {"title": title}
    if template_id is not None:
        payload["linked_template_id"] = template_id
    response = await client.post(
        "/api/requests",
        headers=_headers(user),
        json=payload,
    )
    assert response.status_code == 201, response.text
    return response.json()


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------


async def test_convert_creates_contract_and_links_back(
    client: httpx.AsyncClient, db_session: AsyncSession
) -> None:
    user_org = await _create_user_org(db_session)
    template_id = await _create_template_with_docx(client, user_org.user)
    await _add_required_text_var(
        client, user_org.user, template_id, key="counterparty_name"
    )
    await _add_required_text_var(
        client,
        user_org.user,
        template_id,
        key="effective_date",
        variable_type="date",
        sort_order=1,
    )
    request_row = await _create_request_with_template(
        client, user_org.user, template_id=template_id
    )

    response = await client.post(
        f"/api/requests/{request_row['id']}/convert-to-contract",
        headers=_headers(user_org.user),
        json={
            "title": "Acme NDA",
            "variable_values": {
                "counterparty_name": "Acme Inc.",
                "effective_date": "2026-05-08",
            },
        },
    )
    assert response.status_code == 201, response.text
    body = response.json()

    assert body["contract"]["title"] == "Acme NDA"
    assert body["contract"]["mime_type"] == _DOCX_MIME
    assert body["artifact"]["artifact_type"] == "generated_docx"
    assert body["artifact"]["source"] == "template_generation"
    assert sorted(body["variables_used"]) == [
        "counterparty_name",
        "effective_date",
    ]

    # Request was linked, completed, and inbox item resolved — all in
    # the same transaction.
    assert body["request"]["status"] == "completed"
    assert body["request"]["linked_contract_id"] == body["contract"]["id"]

    # Verify in DB too: a separate read confirms the row actually
    # persisted, not just that we returned a successful response shape.
    request_id = uuid.UUID(request_row["id"])
    request = (
        await db_session.execute(
            select(ContractRequest).where(ContractRequest.id == request_id)
        )
    ).scalar_one()
    assert request.status == "completed"
    assert str(request.linked_contract_id) == body["contract"]["id"]

    inbox_items = (
        await db_session.execute(
            select(InboxItem).where(InboxItem.request_id == request_id)
        )
    ).scalars().all()
    assert len(inbox_items) == 1
    assert inbox_items[0].status == "completed"
    assert inbox_items[0].item_type == "request_review"


async def test_convert_creates_generated_docx_artifact(
    client: httpx.AsyncClient, db_session: AsyncSession
) -> None:
    user_org = await _create_user_org(db_session)
    template_id = await _create_template_with_docx(client, user_org.user)
    await _add_required_text_var(
        client, user_org.user, template_id, key="counterparty_name"
    )
    request_row = await _create_request_with_template(
        client, user_org.user, template_id=template_id
    )

    response = await client.post(
        f"/api/requests/{request_row['id']}/convert-to-contract",
        headers=_headers(user_org.user),
        json={"variable_values": {"counterparty_name": "Acme Inc."}},
    )
    assert response.status_code == 201

    contract_id = uuid.UUID(response.json()["contract"]["id"])
    artifacts = (
        await db_session.execute(
            select(ContractArtifact).where(
                ContractArtifact.contract_id == contract_id
            )
        )
    ).scalars().all()
    assert len(artifacts) == 1
    assert artifacts[0].artifact_type == "generated_docx"
    # Privacy: keys are recorded; values are not.
    assert "counterparty_name" in (
        artifacts[0].metadata_json or {}
    ).get("variable_keys", [])
    assert "variable_values" not in (artifacts[0].metadata_json or {})


async def test_convert_creates_markdown_snapshot_when_conversion_succeeds(
    client: httpx.AsyncClient, db_session: AsyncSession
) -> None:
    user_org = await _create_user_org(db_session)
    template_id = await _create_template_with_docx(client, user_org.user)
    await _add_required_text_var(
        client, user_org.user, template_id, key="counterparty_name"
    )
    request_row = await _create_request_with_template(
        client, user_org.user, template_id=template_id
    )

    response = await client.post(
        f"/api/requests/{request_row['id']}/convert-to-contract",
        headers=_headers(user_org.user),
        json={"variable_values": {"counterparty_name": "Acme"}},
    )
    assert response.status_code == 201
    body = response.json()
    assert body["markdown_snapshot"] is not None
    assert body["markdown_snapshot"]["conversion_status"] == "ready"


async def test_convert_request_without_linked_template_returns_409(
    client: httpx.AsyncClient, db_session: AsyncSession
) -> None:
    user_org = await _create_user_org(db_session)
    user_id = user_org.user.id
    request_row = await _create_request_with_template(
        client, user_org.user, template_id=None
    )

    response = await client.post(
        f"/api/requests/{request_row['id']}/convert-to-contract",
        headers=_headers(user_id),
        json={"variable_values": {}},
    )
    assert response.status_code == 409
    assert "template" in response.json()["detail"].lower()

    # Request must NOT have been mutated.
    fetched = await client.get(
        f"/api/requests/{request_row['id']}",
        headers=_headers(user_id),
    )
    assert fetched.json()["status"] == "open"
    assert fetched.json()["linked_contract_id"] is None


async def test_convert_cancelled_request_returns_409(
    client: httpx.AsyncClient, db_session: AsyncSession
) -> None:
    user_org = await _create_user_org(db_session)
    user_id = user_org.user.id
    template_id = await _create_template_with_docx(client, user_org.user)
    await _add_required_text_var(
        client, user_org.user, template_id, key="counterparty_name"
    )
    request_row = await _create_request_with_template(
        client, user_org.user, template_id=template_id
    )
    cancel = await client.delete(
        f"/api/requests/{request_row['id']}",
        headers=_headers(user_id),
    )
    assert cancel.status_code == 204

    response = await client.post(
        f"/api/requests/{request_row['id']}/convert-to-contract",
        headers=_headers(user_id),
        json={"variable_values": {"counterparty_name": "Acme"}},
    )
    assert response.status_code == 409
    assert "cancelled" in response.json()["detail"].lower()


async def test_convert_already_converted_request_returns_409(
    client: httpx.AsyncClient, db_session: AsyncSession
) -> None:
    user_org = await _create_user_org(db_session)
    user_id = user_org.user.id
    template_id = await _create_template_with_docx(client, user_org.user)
    await _add_required_text_var(
        client, user_org.user, template_id, key="counterparty_name"
    )
    request_row = await _create_request_with_template(
        client, user_org.user, template_id=template_id
    )

    first = await client.post(
        f"/api/requests/{request_row['id']}/convert-to-contract",
        headers=_headers(user_id),
        json={"variable_values": {"counterparty_name": "Acme"}},
    )
    assert first.status_code == 201

    second = await client.post(
        f"/api/requests/{request_row['id']}/convert-to-contract",
        headers=_headers(user_id),
        json={"variable_values": {"counterparty_name": "Acme"}},
    )
    assert second.status_code == 409
    assert "already" in second.json()["detail"].lower()


async def test_convert_propagates_missing_required_variable(
    client: httpx.AsyncClient, db_session: AsyncSession
) -> None:
    user_org = await _create_user_org(db_session)
    template_id = await _create_template_with_docx(client, user_org.user)
    await _add_required_text_var(
        client, user_org.user, template_id, key="counterparty_name"
    )
    request_row = await _create_request_with_template(
        client, user_org.user, template_id=template_id
    )

    response = await client.post(
        f"/api/requests/{request_row['id']}/convert-to-contract",
        headers=_headers(user_org.user),
        json={"variable_values": {}},
    )
    assert response.status_code == 400
    assert "counterparty_name" in response.json()["detail"]


async def test_convert_propagates_unknown_variable(
    client: httpx.AsyncClient, db_session: AsyncSession
) -> None:
    user_org = await _create_user_org(db_session)
    template_id = await _create_template_with_docx(client, user_org.user)
    await _add_required_text_var(
        client, user_org.user, template_id, key="counterparty_name"
    )
    request_row = await _create_request_with_template(
        client, user_org.user, template_id=template_id
    )

    response = await client.post(
        f"/api/requests/{request_row['id']}/convert-to-contract",
        headers=_headers(user_org.user),
        json={
            "variable_values": {
                "counterparty_name": "Acme",
                "rogue": "x",
            }
        },
    )
    assert response.status_code == 400
    assert "rogue" in response.json()["detail"]


async def test_convert_cross_org_returns_404(
    client: httpx.AsyncClient, db_session: AsyncSession
) -> None:
    org_a = await _create_user_org(db_session, email="a@example.com")
    org_b = await _create_user_org(db_session, email="b@example.com")
    template_id = await _create_template_with_docx(client, org_a.user)
    await _add_required_text_var(
        client, org_a.user, template_id, key="counterparty_name"
    )
    request_row = await _create_request_with_template(
        client, org_a.user, template_id=template_id
    )

    response = await client.post(
        f"/api/requests/{request_row['id']}/convert-to-contract",
        headers=_headers(org_b.user),
        json={"variable_values": {"counterparty_name": "Acme"}},
    )
    assert response.status_code == 404


async def test_convert_response_does_not_leak_storage_internals(
    client: httpx.AsyncClient, db_session: AsyncSession
) -> None:
    user_org = await _create_user_org(db_session)
    template_id = await _create_template_with_docx(client, user_org.user)
    await _add_required_text_var(
        client, user_org.user, template_id, key="counterparty_name"
    )
    request_row = await _create_request_with_template(
        client, user_org.user, template_id=template_id
    )

    response = await client.post(
        f"/api/requests/{request_row['id']}/convert-to-contract",
        headers=_headers(user_org.user),
        json={"variable_values": {"counterparty_name": "Acme"}},
    )
    assert response.status_code == 201
    text = response.text
    assert "storage_key" not in text
    assert "wrapped_dek" not in text
    body = response.json()
    assert "storage_key" not in body["artifact"]
    assert "wrapped_dek" not in body["artifact"]
    assert "storage_key" not in body["contract"]


async def test_convert_failure_leaves_request_and_inbox_unchanged(
    client: httpx.AsyncClient,
    db_session: AsyncSession,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """If template generation raises, the request must NOT be marked
    completed and the inbox item must remain open. The endpoint runs
    inside a request-scoped transaction so a service-layer failure
    rolls back any state we touched after the failure point — but more
    importantly, we mutate the request *after* generation succeeds, so
    the assertions below pin both behaviors at once."""
    user_org = await _create_user_org(db_session)
    template_id = await _create_template_with_docx(client, user_org.user)
    await _add_required_text_var(
        client, user_org.user, template_id, key="counterparty_name"
    )
    request_row = await _create_request_with_template(
        client, user_org.user, template_id=template_id
    )

    async def _boom(*_args: Any, **_kwargs: Any) -> Any:
        raise RuntimeError("simulated generation failure")

    monkeypatch.setattr(requests_api, "generate_docx_from_template", _boom)

    with pytest.raises(RuntimeError, match="simulated generation failure"):
        await client.post(
            f"/api/requests/{request_row['id']}/convert-to-contract",
            headers=_headers(user_org.user),
            json={"variable_values": {"counterparty_name": "Acme"}},
        )

    # ``override_get_db`` rolled the session back; reset our test view
    # of it before reading.
    await db_session.rollback()
    request = (
        await db_session.execute(
            select(ContractRequest).where(
                ContractRequest.id == uuid.UUID(request_row["id"])
            )
        )
    ).scalar_one()
    assert request.status == "open"
    assert request.linked_contract_id is None

    inbox = (
        await db_session.execute(
            select(InboxItem).where(
                InboxItem.request_id == uuid.UUID(request_row["id"])
            )
        )
    ).scalars().all()
    assert len(inbox) == 1
    assert inbox[0].status == "open"


# Note: there is no test for "linked template id no longer resolves"
# because the contract_requests.linked_template_id FK to
# agreement_templates makes that state unreachable from a healthy DB.
# The same-org check inside the endpoint is still a defense in depth
# against future schema changes that loosen the FK; for now the FK
# itself is the load-bearing guarantee.
