"""API tests for the agreement template DOCX generation endpoint."""
from __future__ import annotations

import io
import secrets
import subprocess
import uuid
import zipfile
from collections.abc import AsyncIterator, Iterator
from dataclasses import dataclass
from typing import Any

import httpx
import pytest
from docx import Document  # type: ignore[import-not-found]
from docxtpl import DocxTemplate  # type: ignore[import-not-found]
from sqlalchemy import event, select
from sqlalchemy.ext.asyncio import (
    AsyncEngine,
    AsyncSession,
    async_sessionmaker,
    create_async_engine,
)

try:
    from testcontainers.postgres import PostgresContainer
except ImportError:  # pragma: no cover
    PostgresContainer = None  # type: ignore[assignment,misc]

from app.api import agreement_templates as agreement_templates_api  # noqa: E402
from app.core.database import Base, get_db  # noqa: E402
from app.main import app  # noqa: E402
from app.models import (  # noqa: E402
    AgreementTemplate,
    AgreementTemplateArtifact,
    AgreementTemplateMarkdownSnapshot,
    AgreementTemplateVariable,
    Contract,
    ContractArtifact,
    ContractMarkdownSnapshot,
    Organization,
    User,
)
from app.security.audit_log import AuditEvent  # noqa: E402
from app.security.encryption import create_org_master_key  # noqa: E402
from app.services import template_generation as template_generation_service  # noqa: E402
from app.services.document_markdown import MarkdownConversionResult  # noqa: E402
from app.services.storage import StoredDocument  # noqa: E402

_PG_IMAGE = "pgvector/pgvector:pg16"
_INSTANCE_KEY = secrets.token_bytes(32)
_DOCX_MIME = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


def _docker_available() -> bool:
    if PostgresContainer is None:
        return False
    try:
        result = subprocess.run(
            ["docker", "info"], capture_output=True, timeout=5, check=False
        )
    except (FileNotFoundError, subprocess.TimeoutExpired):
        return False
    return result.returncode == 0


def _container_async_url(container: Any) -> str:
    sync_url = container.get_connection_url()
    if sync_url.startswith("postgresql+psycopg2://"):
        return sync_url.replace("postgresql+psycopg2://", "postgresql+asyncpg://", 1)
    if sync_url.startswith("postgresql://"):
        return sync_url.replace("postgresql://", "postgresql+asyncpg://", 1)
    return sync_url


@pytest.fixture(scope="module")
def postgres_container() -> Iterator[Any | None]:
    if not _docker_available() or PostgresContainer is None:
        yield None
        return
    container = PostgresContainer(_PG_IMAGE)
    container.start()
    try:
        yield container
    finally:
        container.stop()


@pytest.fixture
async def engine(postgres_container: Any | None) -> AsyncIterator[AsyncEngine]:
    if postgres_container is None:
        engine = create_async_engine("sqlite+aiosqlite:///:memory:", echo=False)
        tables = [
            Organization.__table__,
            User.__table__,
            AuditEvent.__table__,
            Contract.__table__,
            ContractArtifact.__table__,
            ContractMarkdownSnapshot.__table__,
            AgreementTemplate.__table__,
            AgreementTemplateArtifact.__table__,
            AgreementTemplateMarkdownSnapshot.__table__,
            AgreementTemplateVariable.__table__,
        ]
    else:
        engine = create_async_engine(
            _container_async_url(postgres_container), echo=False
        )
        tables = list(Base.metadata.sorted_tables)

    if engine.dialect.name == "sqlite":
        from pgvector.sqlalchemy import Vector
        from sqlalchemy.ext.compiler import compiles

        @compiles(Vector, "sqlite")
        def _compile_vector_for_sqlite(_t: Any, _c: Any, **_kw: Any) -> str:
            return "BLOB"

        @event.listens_for(engine.sync_engine, "connect")
        def _enable_sqlite_fk(dbapi_connection: Any, _r: Any) -> None:
            cursor = dbapi_connection.cursor()
            cursor.execute("PRAGMA foreign_keys=ON")
            cursor.close()

    async with engine.begin() as conn:
        if engine.dialect.name == "postgresql":
            await conn.exec_driver_sql("CREATE EXTENSION IF NOT EXISTS vector")
        await conn.run_sync(Base.metadata.drop_all, tables=tables)
        await conn.run_sync(Base.metadata.create_all, tables=tables)
    try:
        yield engine
    finally:
        async with engine.begin() as conn:
            await conn.run_sync(Base.metadata.drop_all, tables=tables)
        await engine.dispose()


@pytest.fixture
async def db_session(engine: AsyncEngine) -> AsyncIterator[AsyncSession]:
    maker = async_sessionmaker(engine, expire_on_commit=False, autoflush=False)
    async with maker() as session:
        yield session


@dataclass
class UserOrg:
    org: Organization
    user: User


def _wrapped_org_key(org_id: uuid.UUID) -> bytes:
    return create_org_master_key(
        organization_id=str(org_id),
        instance_key=_INSTANCE_KEY,
    ).to_bytes()


async def _create_user_org(
    session: AsyncSession,
    *,
    email: str | None = None,
) -> UserOrg:
    org = Organization(
        id=uuid.uuid4(),
        name=f"Org {uuid.uuid4()}",
        wrapped_master_key=_wrapped_org_key(uuid.uuid4()),
    )
    org.wrapped_master_key = _wrapped_org_key(org.id)
    user = User(
        id=uuid.uuid4(),
        organization_id=org.id,
        email=email or f"{uuid.uuid4()}@example.com",
        password_hash="hash",
        display_name="Test User",
        is_active=True,
    )
    session.add_all([org, user])
    await session.commit()
    return UserOrg(org=org, user=user)


def _headers(user: User) -> dict[str, str]:
    return {"X-Whereas-Dev-User": str(user.id)}


def _make_nda_template_docx() -> bytes:
    """A small DOCX with two docxtpl placeholders. Bytes are deterministic
    enough for hash assertions in this test, but we don't rely on that."""
    doc = Document()
    doc.add_heading("Mutual NDA", level=1)
    doc.add_paragraph(
        "This Mutual Non-Disclosure Agreement is entered into by "
        "{{counterparty_name}} as of {{effective_date}}."
    )
    doc.add_paragraph(
        "The Parties agree to protect each other's Confidential Information."
    )
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


_NDA_BYTES = _make_nda_template_docx()


def _docx_paragraph_text(blob: bytes) -> list[str]:
    """Return the visible paragraph text of a DOCX.

    docxtpl renders into the same DOCX zip; we open it with python-docx
    on the read side so the assertion is independent of the renderer.
    """
    doc = Document(io.BytesIO(blob))
    return [p.text for p in doc.paragraphs]


class FakeStorage:
    """In-memory storage stub used across upload + retrieve.

    Real ``DocumentStorage`` encrypts before writing to S3. The tests
    don't exercise that path; we just need round-trip semantics so the
    generation service can fetch the template DOCX it just stored.
    """

    _blobs: dict[str, bytes] = {}

    def __init__(self, _settings: Any) -> None:
        pass

    @classmethod
    def reset(cls) -> None:
        cls._blobs.clear()

    async def store_encrypted(
        self,
        *,
        plaintext_bytes: bytes,
        document_id: str,
        org_master_key: bytes,
    ) -> StoredDocument:
        s3_key = f"documents/{document_id}.enc"
        FakeStorage._blobs[s3_key] = plaintext_bytes
        return StoredDocument(
            s3_key=s3_key,
            wrapped_dek_bytes=b"wrapped-dek-" + document_id.encode()[:16],
            encrypted_blob_sha256="a" * 64,
            size_bytes=len(plaintext_bytes),
        )

    async def retrieve_decrypted(
        self,
        *,
        s3_key: str,
        document_id: str,
        wrapped_dek_bytes: bytes,
        org_master_key: bytes,
        expected_blob_sha256: str | None = None,
    ) -> bytes:
        return FakeStorage._blobs[s3_key]


@pytest.fixture
async def client(
    db_session: AsyncSession, monkeypatch: pytest.MonkeyPatch
) -> AsyncIterator[httpx.AsyncClient]:
    monkeypatch.setenv("WHEREAS_INSTANCE_KEY", _INSTANCE_KEY.hex())
    FakeStorage.reset()

    async def override_get_db() -> AsyncIterator[AsyncSession]:
        try:
            yield db_session
            await db_session.commit()
        except Exception:
            await db_session.rollback()
            raise

    app.dependency_overrides[get_db] = override_get_db

    monkeypatch.setattr(agreement_templates_api, "DocumentStorage", FakeStorage)
    # The contract download endpoint instantiates DocumentStorage from
    # its own module, so the round-trip download test needs the same
    # in-memory backing store.
    from app.api import contracts as contracts_api

    monkeypatch.setattr(contracts_api, "DocumentStorage", FakeStorage)

    def _ok_convert(
        *,
        file_bytes: bytes,
        mime_type: str,
        filename: str | None,
        fallback_plain_text: str | None,
    ) -> MarkdownConversionResult:
        return MarkdownConversionResult(
            status="ready",
            markdown_text="# Generated\n\nbody",
            converter_name="fake",
            converter_version="0.0.1",
            warnings=[],
        )

    monkeypatch.setattr(
        agreement_templates_api, "convert_document_to_markdown", _ok_convert
    )
    # The template_generation service routes through the document_markdown
    # helper; patch the seam used inside the helper too.
    from app.services import document_markdown

    monkeypatch.setattr(
        document_markdown, "convert_document_to_markdown", _ok_convert
    )

    class _StubParsed:
        full_text = "Template plain text fallback"

    monkeypatch.setattr(
        agreement_templates_api,
        "parse_document",
        lambda file_bytes, filename: _StubParsed(),
    )

    transport = httpx.ASGITransport(app=app)
    async with httpx.AsyncClient(transport=transport, base_url="http://test") as c:
        yield c
    app.dependency_overrides.clear()


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


async def _create_template_with_docx(
    client: httpx.AsyncClient,
    user: User,
    *,
    docx_bytes: bytes = _NDA_BYTES,
    name: str = "Mutual NDA",
) -> str:
    created = await client.post(
        "/api/agreement-templates",
        headers=_headers(user),
        json={"name": name, "template_type": "NDA"},
    )
    assert created.status_code == 201, created.text
    template_id = created.json()["id"]
    upload = await client.post(
        f"/api/agreement-templates/{template_id}/upload",
        headers=_headers(user),
        files={"file": ("nda.docx", docx_bytes, _DOCX_MIME)},
    )
    assert upload.status_code == 201, upload.text
    return template_id


async def _add_required_text_var(
    client: httpx.AsyncClient,
    user: User,
    template_id: str,
    *,
    key: str,
    label: str | None = None,
    variable_type: str = "text",
    sort_order: int = 0,
) -> None:
    response = await client.post(
        f"/api/agreement-templates/{template_id}/variables",
        headers=_headers(user),
        json={
            "key": key,
            "label": label or key.replace("_", " ").title(),
            "variable_type": variable_type,
            "required": True,
            "sort_order": sort_order,
        },
    )
    assert response.status_code == 201


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------


async def test_generate_creates_contract_and_generated_artifact(
    client: httpx.AsyncClient, db_session: AsyncSession
) -> None:
    user_org = await _create_user_org(db_session)
    template_id = await _create_template_with_docx(client, user_org.user)
    await _add_required_text_var(
        client, user_org.user, template_id, key="counterparty_name"
    )
    await _add_required_text_var(
        client,
        user_org.user,
        template_id,
        key="effective_date",
        variable_type="date",
        sort_order=1,
    )

    response = await client.post(
        f"/api/agreement-templates/{template_id}/generate",
        headers=_headers(user_org.user),
        json={
            "title": "Acme NDA",
            "variable_values": {
                "counterparty_name": "Acme Inc.",
                "effective_date": "2026-05-08",
            },
        },
    )
    assert response.status_code == 201, response.text
    body = response.json()

    assert body["contract"]["title"] == "Acme NDA"
    assert body["contract"]["mime_type"] == _DOCX_MIME
    assert body["artifact"]["artifact_type"] == "generated_docx"
    assert body["artifact"]["source"] == "template_generation"
    assert body["artifact"]["is_official"] is True
    assert sorted(body["variables_used"]) == [
        "counterparty_name",
        "effective_date",
    ]

    # Contract row exists with the generated_docx artifact.
    contract_id = uuid.UUID(body["contract"]["id"])
    contract = (
        await db_session.execute(
            select(Contract).where(Contract.id == contract_id)
        )
    ).scalar_one()
    assert contract.title == "Acme NDA"
    assert contract.s3_key.startswith("documents/")
    artifacts = (
        await db_session.execute(
            select(ContractArtifact).where(
                ContractArtifact.contract_id == contract_id
            )
        )
    ).scalars().all()
    assert len(artifacts) == 1
    assert artifacts[0].artifact_type == "generated_docx"
    assert artifacts[0].metadata_json["template_id"] == template_id
    assert artifacts[0].metadata_json["template_name"] == "Mutual NDA"
    # Privacy: variable keys are recorded, plaintext values are NOT.
    assert sorted(artifacts[0].metadata_json["variable_keys"]) == [
        "counterparty_name",
        "effective_date",
    ]
    assert "variable_values" not in artifacts[0].metadata_json

    # Original template artifact is untouched (still present, still
    # original_upload).
    template_artifacts = (
        await db_session.execute(
            select(AgreementTemplateArtifact).where(
                AgreementTemplateArtifact.template_id == uuid.UUID(template_id)
            )
        )
    ).scalars().all()
    assert {a.artifact_type for a in template_artifacts} == {"original_upload"}


async def test_generated_docx_contains_substituted_values(
    client: httpx.AsyncClient, db_session: AsyncSession
) -> None:
    user_org = await _create_user_org(db_session)
    template_id = await _create_template_with_docx(client, user_org.user)
    await _add_required_text_var(
        client, user_org.user, template_id, key="counterparty_name"
    )
    await _add_required_text_var(
        client,
        user_org.user,
        template_id,
        key="effective_date",
        variable_type="date",
        sort_order=1,
    )

    response = await client.post(
        f"/api/agreement-templates/{template_id}/generate",
        headers=_headers(user_org.user),
        json={
            "variable_values": {
                "counterparty_name": "Acme Inc.",
                "effective_date": "2026-05-08",
            },
        },
    )
    assert response.status_code == 201

    artifact_row = (
        await db_session.execute(
            select(ContractArtifact).where(
                ContractArtifact.contract_id
                == uuid.UUID(response.json()["contract"]["id"])
            )
        )
    ).scalar_one()
    blob = FakeStorage._blobs[artifact_row.storage_key]
    text = "\n".join(_docx_paragraph_text(blob))
    assert "Acme Inc." in text
    assert "2026-05-08" in text
    assert "{{counterparty_name}}" not in text
    assert "{{effective_date}}" not in text


async def test_required_variable_missing_returns_400(
    client: httpx.AsyncClient, db_session: AsyncSession
) -> None:
    user_org = await _create_user_org(db_session)
    template_id = await _create_template_with_docx(client, user_org.user)
    await _add_required_text_var(
        client, user_org.user, template_id, key="counterparty_name"
    )

    response = await client.post(
        f"/api/agreement-templates/{template_id}/generate",
        headers=_headers(user_org.user),
        json={"variable_values": {}},
    )
    assert response.status_code == 400
    assert "counterparty_name" in response.json()["detail"]


async def test_unknown_variable_rejected(
    client: httpx.AsyncClient, db_session: AsyncSession
) -> None:
    user_org = await _create_user_org(db_session)
    template_id = await _create_template_with_docx(client, user_org.user)
    await _add_required_text_var(
        client, user_org.user, template_id, key="counterparty_name"
    )

    response = await client.post(
        f"/api/agreement-templates/{template_id}/generate",
        headers=_headers(user_org.user),
        json={
            "variable_values": {
                "counterparty_name": "Acme",
                "rogue_variable": "x",
            },
        },
    )
    assert response.status_code == 400
    assert "rogue_variable" in response.json()["detail"]


async def test_invalid_date_format_returns_400(
    client: httpx.AsyncClient, db_session: AsyncSession
) -> None:
    user_org = await _create_user_org(db_session)
    template_id = await _create_template_with_docx(client, user_org.user)
    await _add_required_text_var(
        client,
        user_org.user,
        template_id,
        key="effective_date",
        variable_type="date",
    )

    response = await client.post(
        f"/api/agreement-templates/{template_id}/generate",
        headers=_headers(user_org.user),
        json={"variable_values": {"effective_date": "5/8/2026"}},
    )
    assert response.status_code == 400


async def test_cross_org_generation_returns_404(
    client: httpx.AsyncClient, db_session: AsyncSession
) -> None:
    org_a = await _create_user_org(db_session, email="a@example.com")
    org_b = await _create_user_org(db_session, email="b@example.com")
    template_id = await _create_template_with_docx(client, org_a.user)
    await _add_required_text_var(
        client, org_a.user, template_id, key="counterparty_name"
    )

    response = await client.post(
        f"/api/agreement-templates/{template_id}/generate",
        headers=_headers(org_b.user),
        json={"variable_values": {"counterparty_name": "Acme"}},
    )
    assert response.status_code == 404


async def test_generate_without_original_upload_returns_409(
    client: httpx.AsyncClient, db_session: AsyncSession
) -> None:
    user_org = await _create_user_org(db_session)
    created = await client.post(
        "/api/agreement-templates",
        headers=_headers(user_org.user),
        json={"name": "Naked NDA"},
    )
    template_id = created.json()["id"]

    response = await client.post(
        f"/api/agreement-templates/{template_id}/generate",
        headers=_headers(user_org.user),
        json={"variable_values": {}},
    )
    assert response.status_code == 409


async def test_non_docx_source_template_returns_400(
    client: httpx.AsyncClient, db_session: AsyncSession
) -> None:
    user_org = await _create_user_org(db_session)
    # Create a template, then upload a PDF original (PDFs are accepted by
    # the upload endpoint but generation rejects them).
    created = await client.post(
        "/api/agreement-templates",
        headers=_headers(user_org.user),
        json={"name": "PDF-only NDA"},
    )
    template_id = created.json()["id"]
    pdf_bytes = b"%PDF-1.7\n% Whereas synthetic test PDF\n"
    upload = await client.post(
        f"/api/agreement-templates/{template_id}/upload",
        headers=_headers(user_org.user),
        files={"file": ("nda.pdf", pdf_bytes, "application/pdf")},
    )
    assert upload.status_code == 201

    response = await client.post(
        f"/api/agreement-templates/{template_id}/generate",
        headers=_headers(user_org.user),
        json={"variable_values": {}},
    )
    assert response.status_code == 400


async def test_generated_artifact_response_does_not_expose_storage_key(
    client: httpx.AsyncClient, db_session: AsyncSession
) -> None:
    user_org = await _create_user_org(db_session)
    template_id = await _create_template_with_docx(client, user_org.user)
    await _add_required_text_var(
        client, user_org.user, template_id, key="counterparty_name"
    )

    response = await client.post(
        f"/api/agreement-templates/{template_id}/generate",
        headers=_headers(user_org.user),
        json={"variable_values": {"counterparty_name": "Acme"}},
    )
    assert response.status_code == 201
    text = response.text
    assert "storage_key" not in text
    assert "wrapped_dek" not in text
    body = response.json()
    assert "storage_key" not in body["artifact"]
    assert "wrapped_dek" not in body["artifact"]
    assert "storage_key" not in body["contract"]


async def test_generated_artifact_metadata_does_not_persist_variable_values(
    client: httpx.AsyncClient, db_session: AsyncSession
) -> None:
    """Plaintext variable values must not be persisted alongside the
    artifact row — they are already in the rendered DOCX and may
    contain sensitive contract data (counterparty names, dates,
    dollar amounts). The keys are kept for audit; values are not."""
    user_org = await _create_user_org(db_session)
    template_id = await _create_template_with_docx(client, user_org.user)
    await _add_required_text_var(
        client, user_org.user, template_id, key="counterparty_name"
    )

    response = await client.post(
        f"/api/agreement-templates/{template_id}/generate",
        headers=_headers(user_org.user),
        json={
            "variable_values": {
                "counterparty_name": "Acme Health Inc. (HIPAA-covered entity)",
            },
        },
    )
    assert response.status_code == 201

    # Search the response and the persisted metadata row for the
    # value. The string is unique enough that any leak would surface.
    sentinel = "Acme Health Inc."
    assert sentinel not in response.text
    artifact_row = (
        await db_session.execute(
            select(ContractArtifact).where(
                ContractArtifact.contract_id
                == uuid.UUID(response.json()["contract"]["id"])
            )
        )
    ).scalar_one()
    metadata_str = "" if artifact_row.metadata_json is None else str(
        artifact_row.metadata_json
    )
    assert sentinel not in metadata_str
    assert "variable_values" not in (artifact_row.metadata_json or {})
    # ...but the keys ARE recorded for audit.
    assert "counterparty_name" in (artifact_row.metadata_json or {}).get(
        "variable_keys", []
    )


async def test_generated_contract_is_downloadable_via_contract_endpoint(
    client: httpx.AsyncClient, db_session: AsyncSession
) -> None:
    """A generated contract has no `original_upload` artifact —
    download resolution must fall through to the `generated_docx`
    artifact (and recover the original filename from it) instead of
    silently leaning on the legacy ``Contract.s3_key`` column."""
    user_org = await _create_user_org(db_session)
    template_id = await _create_template_with_docx(client, user_org.user)
    await _add_required_text_var(
        client, user_org.user, template_id, key="counterparty_name"
    )

    generation = await client.post(
        f"/api/agreement-templates/{template_id}/generate",
        headers=_headers(user_org.user),
        json={
            "title": "Acme NDA",
            "variable_values": {"counterparty_name": "Acme Inc."},
        },
    )
    assert generation.status_code == 201
    contract_id = generation.json()["contract"]["id"]

    # No original_upload exists on the generated contract — only a
    # generated_docx. The download endpoint must still serve it.
    download = await client.get(
        f"/api/contracts/{contract_id}/download",
        headers=_headers(user_org.user),
    )
    assert download.status_code == 200
    assert download.headers["content-type"] == _DOCX_MIME
    disposition = download.headers["content-disposition"]
    # Filename should reflect the generated artifact's filename, which
    # is derived from the contract title.
    assert "Acme_NDA" in disposition
    # Round-trip: the bytes we downloaded must contain the substituted
    # variable value, proving we actually decrypted the right blob.
    paragraph_text = "\n".join(_docx_paragraph_text(download.content))
    assert "Acme Inc." in paragraph_text


async def test_generation_creates_markdown_snapshot_on_success(
    client: httpx.AsyncClient, db_session: AsyncSession
) -> None:
    user_org = await _create_user_org(db_session)
    template_id = await _create_template_with_docx(client, user_org.user)
    await _add_required_text_var(
        client, user_org.user, template_id, key="counterparty_name"
    )

    response = await client.post(
        f"/api/agreement-templates/{template_id}/generate",
        headers=_headers(user_org.user),
        json={"variable_values": {"counterparty_name": "Acme"}},
    )
    assert response.status_code == 201
    body = response.json()
    assert body["markdown_snapshot"] is not None
    assert body["markdown_snapshot"]["conversion_status"] == "ready"

    contract_id = uuid.UUID(body["contract"]["id"])
    snaps = (
        await db_session.execute(
            select(ContractMarkdownSnapshot).where(
                ContractMarkdownSnapshot.contract_id == contract_id
            )
        )
    ).scalars().all()
    assert len(snaps) == 1
    assert snaps[0].source_kind == "generated"


async def test_generation_succeeds_when_markdown_conversion_fails(
    client: httpx.AsyncClient,
    db_session: AsyncSession,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    user_org = await _create_user_org(db_session)
    template_id = await _create_template_with_docx(client, user_org.user)
    await _add_required_text_var(
        client, user_org.user, template_id, key="counterparty_name"
    )

    def _raise(**_kwargs: Any) -> MarkdownConversionResult:
        raise RuntimeError("markdown converter exploded")

    from app.services import document_markdown

    monkeypatch.setattr(
        document_markdown, "convert_document_to_markdown", _raise
    )

    response = await client.post(
        f"/api/agreement-templates/{template_id}/generate",
        headers=_headers(user_org.user),
        json={"variable_values": {"counterparty_name": "Acme"}},
    )
    assert response.status_code == 201
    body = response.json()
    assert body["markdown_snapshot"] is None

    contract_id = uuid.UUID(body["contract"]["id"])
    snaps = (
        await db_session.execute(
            select(ContractMarkdownSnapshot).where(
                ContractMarkdownSnapshot.contract_id == contract_id
            )
        )
    ).scalars().all()
    assert snaps == []


async def test_default_title_falls_back_to_template_name(
    client: httpx.AsyncClient, db_session: AsyncSession
) -> None:
    user_org = await _create_user_org(db_session)
    template_id = await _create_template_with_docx(
        client, user_org.user, name="Mutual NDA"
    )

    response = await client.post(
        f"/api/agreement-templates/{template_id}/generate",
        headers=_headers(user_org.user),
        json={"variable_values": {}},
    )
    assert response.status_code == 201
    title = response.json()["contract"]["title"]
    assert title.startswith("Mutual NDA")
    assert "generated" in title


async def test_generated_docx_is_a_valid_docx_zip(
    client: httpx.AsyncClient, db_session: AsyncSession
) -> None:
    user_org = await _create_user_org(db_session)
    template_id = await _create_template_with_docx(client, user_org.user)

    response = await client.post(
        f"/api/agreement-templates/{template_id}/generate",
        headers=_headers(user_org.user),
        json={"variable_values": {}},
    )
    assert response.status_code == 201

    artifact_row = (
        await db_session.execute(
            select(ContractArtifact).where(
                ContractArtifact.contract_id
                == uuid.UUID(response.json()["contract"]["id"])
            )
        )
    ).scalar_one()
    blob = FakeStorage._blobs[artifact_row.storage_key]
    with zipfile.ZipFile(io.BytesIO(blob)) as archive:
        names = set(archive.namelist())
    assert "[Content_Types].xml" in names
    assert "word/document.xml" in names


async def test_docxtpl_renders_runs_split_across_xml(
    client: httpx.AsyncClient, db_session: AsyncSession
) -> None:
    """Word frequently splits the same logical run into multiple XML
    runs (e.g. as you type formatting). docxtpl handles this; a naive
    string-replace of ``{{counterparty_name}}`` in document.xml would not.

    This test ensures we picked the right rendering library.
    """
    user_org = await _create_user_org(db_session)

    doc = Document()
    p = doc.add_paragraph()
    # Force the placeholder across multiple runs so the test fails if
    # someone swaps the renderer for a naive str.replace.
    p.add_run("{{")
    p.add_run("counterparty_name")
    p.add_run("}}")
    buf = io.BytesIO()
    doc.save(buf)
    template_id = await _create_template_with_docx(
        client, user_org.user, docx_bytes=buf.getvalue()
    )
    await _add_required_text_var(
        client, user_org.user, template_id, key="counterparty_name"
    )

    response = await client.post(
        f"/api/agreement-templates/{template_id}/generate",
        headers=_headers(user_org.user),
        json={"variable_values": {"counterparty_name": "Acme Inc."}},
    )
    assert response.status_code == 201
    artifact_row = (
        await db_session.execute(
            select(ContractArtifact).where(
                ContractArtifact.contract_id
                == uuid.UUID(response.json()["contract"]["id"])
            )
        )
    ).scalar_one()
    blob = FakeStorage._blobs[artifact_row.storage_key]
    text = "\n".join(_docx_paragraph_text(blob))
    assert "Acme Inc." in text


# ---------------------------------------------------------------------------
# Service-level unit tests for the rendering primitive
# ---------------------------------------------------------------------------


def test_render_docx_substitutes_placeholders() -> None:
    rendered = template_generation_service._render_docx(
        _NDA_BYTES,
        {"counterparty_name": "Acme Inc.", "effective_date": "2026-05-08"},
    )
    paragraphs = _docx_paragraph_text(rendered)
    joined = "\n".join(paragraphs)
    assert "Acme Inc." in joined
    assert "2026-05-08" in joined


def test_render_docx_with_docxtpl_unrelated_template_round_trip() -> None:
    """Round-trip via DocxTemplate + python-docx to guard against
    regressions if docxtpl drops compatibility with our placeholder
    syntax."""
    src = io.BytesIO(_NDA_BYTES)
    tpl = DocxTemplate(src)
    tpl.render({"counterparty_name": "X", "effective_date": "2026-05-08"})
    out = io.BytesIO()
    tpl.save(out)
    paragraphs = _docx_paragraph_text(out.getvalue())
    assert any("X" in p for p in paragraphs)
