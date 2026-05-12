"""Render a structured text diff as a downloadable comparison-report DOCX.

PR #90 — first redline export foundation.

Scope (deliberate):

* This is **not** a Word tracked-changes (``w:ins`` / ``w:del``) file.
  Generating a true tracked-changes DOCX from arbitrary text input is
  error-prone — paragraph boundaries, formatting, and table cells all
  complicate things, and a half-broken file is worse than no file.
* What we produce is a self-contained **comparison report** DOCX
  honestly labelled as such. Each unchanged line is plain; each
  removed line is rendered red + strikethrough; each added line is
  rendered green + underlined. A disclaimer at the top tells the
  reader this is a working aid, not an official Word redline.
* The function takes an already-computed :class:`DiffResult` plus a
  small bag of metadata about the two sides and returns a ``bytes``
  payload ready to stream over HTTP. The route layer is responsible
  for org scoping, decryption, extraction, audit, and headers.
* No persistence here. The route does not write a
  ``ContractArtifact`` row; the DOCX bytes are returned to the user
  and forgotten.
"""
from __future__ import annotations

import io
from dataclasses import dataclass
from datetime import UTC, datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from app.services.artifact_compare import DiffBlock, DiffLine, DiffResult

# Colors picked to look reasonable in Word's default theme without
# leaning on theme colors (which can render differently per template).
_REMOVED_COLOR = RGBColor(0xB9, 0x1C, 0x1C)  # ink danger
_ADDED_COLOR = RGBColor(0x15, 0x80, 0x3D)  # ink success
_MUTED_COLOR = RGBColor(0x6B, 0x72, 0x80)  # ink subtle


@dataclass(frozen=True)
class CompareSideMetadata:
    """Display-only metadata about one side of the comparison.

    All fields are safe for inclusion in the DOCX header — the route
    only forwards artifact id, user-facing label, and the user-supplied
    filename. Storage internals and signer PII never reach this layer.
    """

    label: str
    filename: str | None
    created_at: datetime | None


def render_compare_report_docx(
    *,
    diff: DiffResult,
    base: CompareSideMetadata,
    compare: CompareSideMetadata,
    contract_title: str,
) -> bytes:
    """Render ``diff`` as a comparison-report DOCX and return its bytes.

    The document layout is:

    1. Title block: contract title.
    2. Disclaimer paragraph (red, italic) that this is a text
       comparison preview, not an official Word redline.
    3. "Versions" header showing the two sides' label / filename /
       timestamp.
    4. "Summary" block: added / removed / changed counts from the
       diff summary.
    5. "Differences" section: one paragraph per diff block. Inside a
       ``changed`` block we keep the difflib order of "removed lines
       first, then added lines" so the reader can scan the swap.

    Context (unchanged) blocks are collapsed to a single muted "..."
    paragraph to keep the report focused on the actual changes — the
    full text is still available in the original artifacts.
    """
    document = Document()

    # Title.
    title = document.add_heading(level=1)
    run = title.add_run("Comparison report")
    run.bold = True

    subtitle = document.add_paragraph()
    subtitle_run = subtitle.add_run(contract_title)
    subtitle_run.italic = True
    subtitle_run.font.color.rgb = _MUTED_COLOR

    # Disclaimer.
    disclaimer = document.add_paragraph()
    disclaimer_run = disclaimer.add_run(
        "This is a text comparison preview, not an official Word "
        "redline. Differences are highlighted as a working aid; the "
        "underlying official documents remain authoritative."
    )
    disclaimer_run.italic = True
    disclaimer_run.font.color.rgb = _REMOVED_COLOR

    # Versions block.
    document.add_heading("Versions", level=2)
    _add_side_paragraph(document, side_label="Left version", side=base)
    _add_side_paragraph(document, side_label="Right version", side=compare)

    # Summary. PR #93 — the unit is now a paragraph (see
    # ``_split_paragraphs`` in the artifact_compare service); the
    # underlying ``DiffSummary`` field names stay as ``added_lines``
    # etc. for back-compat with persisted ``metadata_json`` on saved
    # redlines, but the user-facing labels in the DOCX use the right
    # word.
    document.add_heading("Summary", level=2)
    summary_para = document.add_paragraph()
    summary_para.add_run(
        f"Added paragraphs: {diff.summary.added_lines}    "
        f"Removed paragraphs: {diff.summary.removed_lines}    "
        f"Changed blocks: {diff.summary.changed_blocks}    "
        f"Unchanged paragraphs: {diff.summary.unchanged_lines}"
    )
    if diff.warnings:
        warning_para = document.add_paragraph()
        warning_run = warning_para.add_run(
            "Warnings: " + ", ".join(sorted(set(diff.warnings)))
        )
        warning_run.italic = True
        warning_run.font.color.rgb = _MUTED_COLOR

    # Differences.
    document.add_heading("Differences", level=2)
    if not diff.diff_blocks:
        empty_para = document.add_paragraph()
        empty_run = empty_para.add_run(
            "No differences in the compared portion of these versions."
        )
        empty_run.italic = True
        empty_run.font.color.rgb = _MUTED_COLOR
    else:
        for block in diff.diff_blocks:
            _add_block(document, block)

    # Footer.
    footer = document.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run(
        f"Generated {datetime.now(tz=UTC).strftime('%Y-%m-%d %H:%M UTC')}"
    )
    footer_run.italic = True
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = _MUTED_COLOR

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def _add_side_paragraph(
    document, *, side_label: str, side: CompareSideMetadata
) -> None:
    para = document.add_paragraph()
    label_run = para.add_run(f"{side_label}: ")
    label_run.bold = True
    para.add_run(side.label)
    if side.filename:
        para.add_run(f" — {side.filename}")
    if side.created_at:
        ts = side.created_at.strftime("%Y-%m-%d")
        muted = para.add_run(f"  (added {ts})")
        muted.italic = True
        muted.font.color.rgb = _MUTED_COLOR


def _add_block(document, block: DiffBlock) -> None:
    """Render one diff block into the DOCX body, paragraph-aware.

    The unit is now a paragraph (see :func:`_split_paragraphs` in the
    artifact_compare service): ``context`` blocks collapse to a muted
    "N unchanged paragraphs" indicator, ``added`` and ``removed``
    blocks get a small section label, and ``changed`` blocks split
    visually into "Before:" (the removed paragraphs) and "After:"
    (the added paragraphs) so a legal reviewer can scan the swap.
    """
    if block.type == "context":
        # Collapse runs of equal text to a single muted indicator so
        # the report stays focused on differences. Users with the
        # original artifacts can still inspect unchanged text directly.
        para = document.add_paragraph()
        muted = para.add_run(
            f"… {len(block.lines)} unchanged paragraph"
            f"{'s' if len(block.lines) != 1 else ''} …"
        )
        muted.italic = True
        muted.font.color.rgb = _MUTED_COLOR
        return

    if block.type == "added":
        _add_section_label(document, "Added paragraphs", tone="added")
        for line in block.lines:
            _add_line(document, line)
        return

    if block.type == "removed":
        _add_section_label(document, "Removed paragraphs", tone="removed")
        for line in block.lines:
            _add_line(document, line)
        return

    # ``changed`` block: split the inner lines into removed-first /
    # added-after (the order difflib hands us) and render them under
    # explicit Before / After labels so a reviewer can scan the swap.
    _add_section_label(document, "Changed paragraph", tone="changed")
    removed_lines = [line for line in block.lines if line.type == "removed"]
    added_lines = [line for line in block.lines if line.type == "added"]
    if removed_lines:
        _add_subsection_label(document, "Before:")
        for line in removed_lines:
            _add_line(document, line)
    if added_lines:
        _add_subsection_label(document, "After:")
        for line in added_lines:
            _add_line(document, line)


def _add_section_label(document, label: str, *, tone: str) -> None:
    """Render a short, tone-coded label paragraph above a diff block.

    Matches the spirit of a legal redline summary: each block has a
    one-line header so the reader knows what changed without reading
    the whole paragraph first.
    """
    para = document.add_paragraph()
    run = para.add_run(label)
    run.bold = True
    if tone == "added":
        run.font.color.rgb = _ADDED_COLOR
    elif tone == "removed":
        run.font.color.rgb = _REMOVED_COLOR
    else:  # changed (or any future tone)
        run.font.color.rgb = _MUTED_COLOR


def _add_subsection_label(document, label: str) -> None:
    """Render a small muted sub-label ("Before:" / "After:")."""
    para = document.add_paragraph()
    run = para.add_run(label)
    run.italic = True
    run.font.color.rgb = _MUTED_COLOR


def _add_line(document, line: DiffLine) -> None:
    para = document.add_paragraph()
    run = para.add_run(line.text or " ")
    if line.type == "removed":
        run.font.strike = True
        run.font.color.rgb = _REMOVED_COLOR
    elif line.type == "added":
        run.font.underline = True
        run.font.color.rgb = _ADDED_COLOR
    # context (rare here — context blocks are short-circuited above)
    # falls through unstyled.


def build_export_filename(contract_title: str) -> str:
    """Build a safe ASCII filename for the comparison-report download.

    Strips path separators and non-ASCII characters, collapses
    whitespace to underscores, and clips to a reasonable length so
    the resulting ``Content-Disposition: attachment; filename=...``
    header stays valid and predictable across browsers.
    """
    safe_chars: list[str] = []
    for ch in (contract_title or "comparison-report").strip():
        if ch.isascii() and (ch.isalnum() or ch in (" ", "-", "_")):
            safe_chars.append(ch)
        else:
            safe_chars.append("_")
    safe = "".join(safe_chars).strip().replace(" ", "_")
    if not safe:
        safe = "comparison-report"
    safe = safe[:80]
    return f"{safe}-comparison-report.docx"
